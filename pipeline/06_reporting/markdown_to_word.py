# -*- coding: utf-8 -*-
"""
Combine multiple Markdown documents in docs/thesis into a single Word file.
"""
from __future__ import annotations

from pathlib import Path
from typing import Iterable, List

from docx import Document


def _preferred_order(files: Iterable[Path]) -> List[Path]:
    priority = [
        "DAN_Y_KHOA_LUAN.md",
        "HUONG_DAN_SU_DUNG_KHOA_LUAN.md",
        "HUONG_DAN_VIET_KHOA_LUAN.md",
        "KHOA_LUAN_CHUONG_1.md",
        "KHOA_LUAN_CHUONG_2.md",
        "KHOA_LUAN_CHUONG_3.md",
        "KHOA_LUAN_CHUONG_4.md",
        "KHOA_LUAN_CHUONG_5.md",
        "KHOA_LUAN_HOAN_CHINH.md",
        "README_KHOA_LUAN.md",
        "TONG_HOP_CUOI_CUNG.md",
        "Y_NGHIA_VA_UNG_DUNG.md",
    ]
    lookup = {p.name: p for p in files}
    ordered = [lookup[name] for name in priority if name in lookup]
    remaining = sorted(
        [p for p in files if p.name not in priority],
        key=lambda p: p.name.lower(),
    )
    return ordered + remaining


def _add_markdown_block(doc: Document, line: str) -> None:
    stripped = line.strip("\n")

    if not stripped:
        doc.add_paragraph("")
        return

    if stripped.startswith("#"):
        level = len(stripped) - len(stripped.lstrip("#"))
        text = stripped[level:].strip()
        level = min(level, 4)
        doc.add_heading(text if text else "", level=level)
        return

    if stripped.startswith(("- ", "* ")):
        doc.add_paragraph(stripped[2:], style="List Bullet")
        return

    if stripped.startswith(">"):
        doc.add_paragraph(stripped[1:].strip(), style="Intense Quote")
        return

    doc.add_paragraph(stripped)


def add_markdown_file(doc: Document, path: Path) -> None:
    doc.add_heading(path.stem.replace("_", " "), level=1)
    doc.add_paragraph(f"(Nguồn: {path.name})", style="Intense Quote")
    with path.open(encoding="utf-8") as f:
        in_code_block = False
        code_buffer: List[str] = []
        for line in f:
            if line.strip().startswith("```"):
                if in_code_block:
                    doc.add_paragraph(
                        "\n".join(code_buffer),
                        style="Intense Quote",
                    )
                    code_buffer = []
                in_code_block = not in_code_block
                continue

            if in_code_block:
                code_buffer.append(line.rstrip("\n"))
            else:
                _add_markdown_block(doc, line)

        if code_buffer:
            doc.add_paragraph("\n".join(code_buffer), style="Intense Quote")

    doc.add_page_break()


def convert_folder_to_docx(folder: Path, output: Path) -> None:
    md_files = list(folder.glob("*.md"))
    if not md_files:
        raise FileNotFoundError(f"No markdown files found in {folder}")

    doc = Document()
    doc.add_heading("KHÓA LUẬN - FILE HỢP NHẤT", 0)
    doc.add_paragraph(
        "Tài liệu này gộp toàn bộ nội dung Markdown trong thư mục docs/thesis "
        "để tiện chỉnh sửa trực tiếp trong Word. Mỗi phần giữ nguyên hướng dẫn "
        "về nơi chèn hình/công thức/bảng.",
    )

    for md_file in _preferred_order(md_files):
        add_markdown_file(doc, md_file)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Saved combined thesis document to {output}")


if __name__ == "__main__":
    project_root = Path(__file__).resolve().parents[2]
    folder = project_root / "docs" / "thesis"
    output_file = folder / "KHOA_LUAN_FULL.docx"
    convert_folder_to_docx(folder, output_file)

