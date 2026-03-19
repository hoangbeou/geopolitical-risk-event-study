"""
Build a Chapter 1 draft directly inside the FTU Word template.

The script reads the existing template (preserving cover, prelim pages, and
later chapter placeholders), removes the placeholder bullets under Chapter 1,
and injects the fully written content so the user can continue editing inside
Word without reformatting everything manually.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
TEMPLATE_PATH = Path("docs/thesis/templates/KLTN_FORMAT_TEMPLATE.docx")
OUTPUT_PATH = Path("docs/thesis/drafts/KLTN_CHAPTER1_DRAFT.docx")


def remove_paragraph(paragraph: Paragraph) -> None:
    """Safely delete a paragraph from the document."""
    element = paragraph._element
    parent = element.getparent()
    if parent is None:
        return
    parent.remove(element)
    paragraph._p = paragraph._element = None


def truncate_after_chapter_one(doc: Document) -> None:
    """
    Keep everything up to the Chapter 1 title and delete the remaining placeholders.

    This leaves the cover, prelim pages, and the Chapter 1 title intact, while
    freeing the remainder so we can append the real content.
    """

    chapter_title_idx = None
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip().upper().startswith("TỔNG QUAN"):
            chapter_title_idx = idx
            break

    if chapter_title_idx is None:
        return

    # Remove paragraphs after the Chapter 1 title
    idx = len(doc.paragraphs) - 1
    while idx > chapter_title_idx:
        remove_paragraph(doc.paragraphs[idx])
        idx -= 1


def add_heading(doc: Document, text: str, level: int = 1) -> Paragraph:
    """Add a bold heading paragraph."""
    para = doc.add_paragraph(text)
    run = para.runs[0]
    run.bold = True
    if level == 1:
        para.paragraph_format.first_line_indent = 0
    return para


def add_body(doc: Document, text: str) -> None:
    """Add a standard body paragraph with FTU formatting."""
    for paragraph in text.strip().split("\n"):
        doc.add_paragraph(paragraph.strip())


CHAPTER_1_CONTENT = [
    {
        "heading": "1.1 Tổng quan nghiên cứu",
        "subsections": [
            {
                "subheading": "1.1.1 Tài sản trú ẩn truyền thống",
                "paragraphs": [
                    (
                        "Nghiên cứu kinh điển của Baur và Lucey (2010) xác lập vàng như tài sản vừa mang "
                        "tính phòng thủ (hedge) dài hạn vừa đóng vai trò trú ẩn tạm thời khi khủng hoảng "
                        "tài chính bùng phát. Các tác giả sau đó như Baur và McDermott (2016) hay Beckmann "
                        "và Czudaj (2017) mở rộng khái niệm này sang bạc, franc Thụy Sĩ và trái phiếu chính "
                        "phủ Mỹ. Kết quả chung chỉ ra vàng phản ứng mạnh với cú sốc chính trị nhờ độ sâu thị "
                        "trường và tính thanh khoản toàn cầu."
                    ),
                    (
                        "Trong giai đoạn hậu khủng hoảng 2008, Basu et al. (2018) và Shahzad et al. (2019) "
                        "cho thấy vàng duy trì khả năng giảm thiểu rủi ro tail risk, song hiệu quả phụ thuộc "
                        "vào thời điểm và khu vực. Các thị trường mới nổi – trong đó có Việt Nam – thường "
                        "chứng kiến phản ứng mạnh hơn vì nhà đầu tư cá nhân chiếm tỷ trọng lớn và ít công cụ "
                        "phòng vệ khác. Điều này giải thích tại sao vàng vẫn được dùng làm chuẩn để benchmark "
                        "tài sản trú ẩn trong khóa luận."
                    ),
                ],
            },
            {
                "subheading": "1.1.2 Bitcoin và cơ chế trú ẩn thay thế",
                "paragraphs": [
                    (
                        "Dyhrberg (2016) và Bouri et al. (2017) là những nghiên cứu đầu tiên xem Bitcoin như "
                        "tài sản mang đặc tính lai giữa vàng và USD. Trong các cú sốc ngắn hạn, Bitcoin có thể "
                        "tăng giá tương tự vàng; tuy nhiên kết quả thiếu nhất quán khi biến động xuất phát từ "
                        "địa chính trị. Corbet et al. (2020) cho thấy Bitcoin chỉ cung cấp phòng vệ hạn chế "
                        "trước căng thẳng Mỹ–Iran, trong khi Ji et al. (2021) ghi nhận cơ chế flight-to-safety "
                        "chỉ tồn tại ở phân vị lợi suất cao."
                    ),
                    (
                        "Kháng cự kiểm soát vốn là yếu tố khiến Bitcoin trở thành điểm đến của dòng tiền khi "
                        "xung đột xảy ra tại các nền kinh tế châu Á. Nghiên cứu của Fang et al. (2022) cho thấy "
                        "lượng stablecoin đổ vào các sàn OTC Trung Quốc tăng mạnh mỗi khi có thông tin siết vốn, "
                        "kéo theo nhu cầu Bitcoin tăng. Những bằng chứng này gợi mở khả năng phản ứng vùng là chìa "
                        "khóa để giải thích vai trò trú ẩn của Bitcoin."
                    ),
                ],
            },
            {
                "subheading": "1.1.3 Đo lường rủi ro địa chính trị và phản ứng đa tài sản",
                "paragraphs": [
                    (
                        "Caldara và Iacoviello (2022) phát triển chỉ số Geopolitical Risk (GPR) dựa trên tần suất "
                        "từ khóa báo chí, đồng thời tách riêng hai thành phần GPRD_ACT (hành động) và GPRD_THREAT "
                        "(đe dọa). Các nghiên cứu nối tiếp như Balcilar et al. (2022) hay Yarovaya et al. (2022) "
                        "kết hợp GPR với thị trường tiền mã hóa nhưng thường chỉ dừng ở mức kiểm định tuyến tính "
                        "hoặc VAR, chưa khai thác cấu trúc ACT/THREAT."
                    ),
                    (
                        "Quan hệ giữa GPR và vàng/dầu đã khá rõ ràng (Antonakakis et al., 2017), song GPR với "
                        "Bitcoin hoặc cặp BTC–GOLD lại tạo ra kết quả mâu thuẫn. Một số nghiên cứu giai đoạn "
                        "COVID-19 cho rằng Bitcoin là trú ẩn, trong khi dữ liệu chiến tranh Nga–Ukraine lại cho "
                        "thấy Bitcoin giảm cùng vàng do nhu cầu thanh khoản USD. Điều đó đòi hỏi cách tiếp cận "
                        "định tính hơn kết hợp đặc điểm sự kiện, vùng và mức độ đe dọa."
                    ),
                ],
            },
            {
                "subheading": "1.1.4 Nhận định tổng quan cho bối cảnh Việt Nam",
                "paragraphs": [
                    (
                        "Các nghiên cứu trong nước chủ yếu tập trung vào tác động của rủi ro chính trị tới "
                        "tỷ giá hoặc thị trường chứng khoán (Nguyễn & Phạm, 2020). Nghiên cứu về Bitcoin "
                        "vẫn còn hạn chế do dữ liệu chính thức khan hiếm. Trong khi đó, nhà đầu tư Việt Nam "
                        "ngày càng tham gia các sàn quốc tế và chịu ảnh hưởng tâm lý từ sự kiện toàn cầu."
                    ),
                    (
                        "Do đó, tổng quan văn liệu cho thấy khoảng trống lớn ở chỗ: (i) thiếu bảng sự kiện "
                        "địa chính trị được gắn ngữ cảnh cụ thể; (ii) thiếu phân tích phối hợp giữa BTC, "
                        "GOLD và OIL trong cùng khuôn khổ; và (iii) chưa xem xét sâu vai trò của thành phần "
                        "ACT/THREAT để giải thích hành vi khác biệt giữa các vùng. Những nhận định này dẫn "
                        "trực tiếp tới thiết kế nghiên cứu của khóa luận."
                    ),
                ],
            },
        ],
    },
    {
        "heading": "1.2 Cơ sở lý thuyết và khung phương pháp",
        "subsections": [
            {
                "subheading": "1.2.1 Lý thuyết tài sản trú ẩn, flight-to-quality và flight-to-liquidity",
                "paragraphs": [
                    (
                        "Lý thuyết portfolio của Markowitz (1952) đặt nền tảng cho việc tìm kiếm tài sản "
                        "giảm tương quan khi biến động tăng. Khi cú sốc địa chính trị xảy ra, nhà đầu tư "
                        "có xu hướng flight-to-quality (chuyển sang tài sản an toàn) và flight-to-liquidity "
                        "(ưu tiên tài sản dễ chuyển đổi). Vàng đáp ứng tốt cả hai tiêu chí nên thường tăng "
                        "giá trong các pha “action” rõ ràng."
                    ),
                    (
                        "Bitcoin, ngược lại, vừa có đặc tính rủi ro cao vừa thiếu cơ chế hỗ trợ chính thức, "
                        "nên phản ứng của nó phụ thuộc vào kỳ vọng dòng vốn tương lai. Nếu sự kiện chỉ dừng "
                        "ở mức đe dọa, nhiều nhà đầu tư coi Bitcoin là kênh thay thế để phòng rủi ro kiểm "
                        "soát vốn, tạo ra hiện tượng flight-to-innovation thay vì flight-to-quality. Khung "
                        "lý thuyết này giải thích tại sao cần phân biệt THREAT và ACT."
                    ),
                ],
            },
            {
                "subheading": "1.2.2 Phương pháp Event Study và đo lường CAR",
                "paragraphs": [
                    (
                        "Event Study sử dụng mô hình lợi suất kỳ vọng (mean-adjusted hoặc market model) "
                        "được ước lượng trong cửa sổ trước sự kiện (estimation window). Abnormal Return "
                        "(AR) tại mỗi ngày là chênh lệch giữa lợi suất thực tế và kỳ vọng; Cumulative "
                        "Abnormal Return (CAR) là tổng AR trong cửa sổ sự kiện. Cách tiếp cận này phù hợp "
                        "để ghi nhận phản ứng tức thì của BTC, GOLD và OIL trước 61 sự kiện được gắn ngày "
                        "rõ ràng."
                    ),
                    (
                        "Khóa luận sử dụng cửa sổ (-10, +10) ngày nhằm dung hòa giữa việc thu đủ phản ứng "
                        "và hạn chế nhiễu từ các tin tức vĩ mô khác. Ngoài việc thống kê CAR cho từng sự kiện, "
                        "nghiên cứu còn phân nhóm theo loại sự kiện (War, Political, Mixed), bối cảnh "
                        "tension/action và vùng địa lý để phục vụ phân tích định tính sau này."
                    ),
                ],
            },
            {
                "subheading": "1.2.3 Lý thuyết wavelet và Quantile-on-Quantile Regression (QQR)",
                "paragraphs": [
                    (
                        "Wavelet MODWT cho phép phân rã chuỗi thời gian thành nhiều thang tần số mà không "
                        "giảm số quan sát, phù hợp với dữ liệu tài chính có tính không dừng. Bộ lọc LA8 "
                        "(triển khai bằng sym4 trong PyWavelets) được lựa chọn để tiệm cận thiết kế trong "
                        "nghiên cứu tham chiếu “Dynamic responses of Bitcoin, gold, and green bonds to "
                        "geopolitical risk” – vốn sử dụng wavelet gần đối xứng để hạn chế méo pha."
                    ),
                    (
                        "QQR (Sim & Zhou, 2015) mở rộng quantile regression bằng cách ánh xạ phân vị của "
                        "biến giải thích vào phân vị của biến phụ thuộc, giúp quan sát phụ thuộc phi tuyến "
                        "ở từng trạng thái thị trường. Khi kết hợp với MODWT, QQR trong khóa luận đánh giá "
                        "tương tác giữa phân vị GPR và phân vị lợi suất BTC/GOLD ở chín thang thời gian, từ "
                        "đó lý giải vì sao hệ số trung bình nhỏ nhưng hệ số cực trị lại có ý nghĩa thống kê."
                    ),
                ],
            },
        ],
    },
    {
        "heading": "1.3 Khoảng trống nghiên cứu và đóng góp dự kiến",
        "subsections": [
            {
                "subheading": "1.3.1 Khoảng trống học thuật",
                "paragraphs": [
                    (
                        "Thứ nhất, chưa có nghiên cứu nào xây dựng tập sự kiện địa chính trị với mô tả "
                        "ngữ cảnh (tension vs. action, start vs. escalation) rồi gắn kết trực tiếp với "
                        "CAR của cả BTC, GOLD và OIL. Các nghiên cứu trước thường sử dụng chỉ báo GPR "
                        "tổng hợp và bỏ qua các khác biệt vùng."
                    ),
                    (
                        "Thứ hai, các nghiên cứu Wavelet QQR hiện hữu chủ yếu tập trung vào giai đoạn "
                        "COVID-19 hoặc so sánh BTC với green bonds. Việc cập nhật giai đoạn 2022–2025 với "
                        "chiến tranh Nga–Ukraine, Israel–Hamas và khủng hoảng eo biển Đài Loan vẫn là khoảng "
                        "trống đáng kể."
                    ),
                    (
                        "Thứ ba, vai trò của thành phần GPRD_ACT và GPRD_THREAT hầu như chưa được tích hợp "
                        "vào phân tích định lượng. Đây là điểm mấu chốt để lý giải vì sao Bitcoin đôi khi "
                        "tăng cùng vàng nhưng đôi khi lại diễn biến ngược chiều."
                    ),
                ],
            },
            {
                "subheading": "1.3.2 Đóng góp của khóa luận",
                "paragraphs": [
                    (
                        "Khóa luận kết hợp Event Study (định lượng) và phân tích định tính (context, vùng, "
                        "bên hưởng lợi) để xây dựng bốn pattern phản ứng BTC–GOLD. Báo cáo cũng cung cấp "
                        "Regime Map, Geopolitical Heatmap và các biểu đồ ACT/THREAT, giúp nhà nghiên cứu "
                        "quan sát trực quan cơ chế flight-to-quality hay flight-to-innovation."
                    ),
                    (
                        "Việc thay đổi bộ lọc wavelet sang LA8 (sym4) và mở rộng chín thang thời gian giúp "
                        "phù hợp hoàn toàn với bài báo tham chiếu, đồng thời tạo nền tảng để so sánh biên "
                        "độ hệ số QQR. Các phát hiện về sự tương đồng số lượng sự kiện BTC↑/BTC↓ và "
                        "GOLD↑/GOLD↓ (32/29) cũng là đóng góp gợi ý cho việc thiết kế danh mục hedging."
                    ),
                    (
                        "Đối với thực tiễn, khóa luận đưa ra khuyến nghị rõ ràng về cách nhà đầu tư Việt Nam "
                        "theo dõi ACT/THREAT và vùng xảy ra sự kiện để điều chỉnh tỷ trọng BTC – GOLD. Đây là "
                        "những hàm ý quan trọng cho luận án và các nghiên cứu tiếp theo trong bối cảnh thị "
                        "trường tiền mã hóa nội địa vẫn thiếu khung phân tích chuẩn hóa."
                    ),
                ],
            },
        ],
    },
]


def write_chapter_one() -> None:
    doc = Document(TEMPLATE_PATH)
    truncate_after_chapter_one(doc)

    for section in CHAPTER_1_CONTENT:
        add_heading(doc, section["heading"], level=1)
        for subsection in section["subsections"]:
            add_heading(doc, subsection["subheading"], level=2)
            for paragraph in subsection["paragraphs"]:
                add_body(doc, paragraph)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    write_chapter_one()

