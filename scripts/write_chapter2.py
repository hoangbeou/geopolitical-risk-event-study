"""
Generate a standalone Word draft for Chapter 2: Data and Methodology.

This script does NOT modify the FTU template. Instead, it creates a
simple, self-contained .docx file containing only Chapter 2 so the
user can review and edit it freely.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


OUTPUT_PATH = Path("docs/thesis/drafts/CHAPTER2_DU_LIEU_VA_PHUONG_PHAP.docx")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a bold heading paragraph."""
    para = doc.add_paragraph(text)
    if para.runs:
        para.runs[0].bold = True
    if level == 1:
        para.paragraph_format.first_line_indent = 0


def add_body(doc: Document, text: str) -> None:
    """Add standard body paragraphs (split by line breaks)."""
    for paragraph in text.strip().split("\n"):
        para_text = paragraph.strip()
        if not para_text:
            continue
        doc.add_paragraph(para_text)


def build_chapter_two() -> None:
    doc = Document()

    # Chapter title
    add_heading(doc, "CHƯƠNG 2: DỮ LIỆU VÀ PHƯƠNG PHÁP NGHIÊN CỨU", level=1)

    intro = (
        "Chương này trình bày khung khổ nghiên cứu, quy trình xử lý dữ liệu và hệ thống các "
        "phương pháp định lượng được sử dụng để đánh giá tác động của rủi ro địa chính trị (GPR) "
        "lên hành động giá của Vàng (GOLD), Bitcoin (BTC) và Dầu thô (OIL). Để trả lời các câu "
        "hỏi nghiên cứu, khóa luận xây dựng một khung phân tích dựa trên hai trụ cột phương pháp "
        "song song và bổ trợ lẫn nhau: (i) Wavelet Quantile-on-Quantile Regression (Wavelet QQR) "
        "nhằm mô tả cấu trúc phụ thuộc phi tuyến và đa thang đo thời gian trên toàn bộ mẫu dữ liệu; "
        "và (ii) Nghiên cứu sự kiện (Event Study) nhằm phân loại các chế độ phản ứng giá xung quanh "
        "các cú sốc địa chính trị cụ thể."
    )
    add_body(doc, intro)

    # 2.1 Khung nghiên cứu và thiết kế phân tích
    add_heading(doc, "2.1. Khung nghiên cứu và thiết kế phân tích", level=1)

    sec_21 = (
        "Dựa trên cơ sở lý thuyết đã thảo luận ở Chương 1, khung nghiên cứu được thiết kế với hai "
        "lớp phân tích nhằm cung cấp cái nhìn toàn diện từ tổng quát đến chi tiết.\n\n"
        "Lớp phân tích cấu trúc (Structural Analysis – Global View): Sử dụng phương pháp Wavelet QQR "
        "để phân rã chuỗi thời gian của GPR và lợi suất tài sản. Mục tiêu là ước lượng ma trận hệ số "
        "β(θ, τ), mô tả độ nhạy cảm của tài sản tại các trạng thái thị trường khác nhau (phân vị θ) "
        "và trên các khung thời gian khác nhau (các thang wavelet). Lớp phân tích này cung cấp “bức "
        "tranh nền” về mối quan hệ phụ thuộc.\n\n"
        "Lớp phân tích sự kiện (Event Analysis – Local View): Sử dụng phương pháp Event Study để "
        "tính toán Lợi suất Bất thường (AR) và Lợi suất Bất thường Tích lũy (CAR) xung quanh 61 sự "
        "kiện địa chính trị trọng yếu. Kết quả này được sử dụng để định danh các mẫu hình phản ứng "
        "giá cụ thể của BTC, GOLD và OIL theo từng vùng địa lý và loại hình sự kiện.\n\n"
        "Sự kết hợp này cho phép khóa luận giải quyết vấn đề từ hai phía: Wavelet QQR giải thích cấu "
        "trúc tiềm ẩn của mối quan hệ, trong khi Event Study cung cấp bằng chứng thực nghiệm cụ thể "
        "để kiểm chứng và giải thích các cấu trúc đó."
    )
    add_body(doc, sec_21)

    # 2.2 Dữ liệu nghiên cứu
    add_heading(doc, "2.2. Dữ liệu nghiên cứu", level=1)

    # 2.2.1 Dữ liệu giá tài sản
    add_heading(doc, "2.2.1. Dữ liệu giá tài sản", level=2)
    sec_221 = (
        "Nghiên cứu sử dụng dữ liệu chuỗi thời gian tần suất ngày (daily) cho ba lớp tài sản đại diện:\n\n"
        "- Bitcoin (BTC): Giá đóng cửa hàng ngày (USD), đại diện cho lớp tài sản số.\n"
        "- Vàng (GOLD): Giá vàng giao ngay hoặc hợp đồng tương lai liền kề (USD/oz), đại diện cho kim "
        "loại quý có vai trò quan trọng trong danh mục đầu tư.\n"
        "- Dầu thô (OIL): Giá dầu Brent (USD/thùng), đại diện cho hàng hóa chiến lược và rủi ro năng lượng.\n\n"
        "Các chuỗi dữ liệu được đồng bộ hóa về thời gian giao dịch, loại bỏ các ngày nghỉ lễ và các quan "
        "sát thiếu để đảm bảo tính nhất quán. Lợi suất của tài sản được tính toán dưới dạng lợi suất "
        "logarit liên tục (log-returns):\n\n"
        "r_{i,t} = ln(P_{i,t}) - ln(P_{i,t-1})\n\n"
        "trong đó P_{i,t} là giá đóng cửa của tài sản i tại ngày t. Việc sử dụng lợi suất logarit giúp ổn "
        "định phương sai chuỗi dữ liệu và thuận lợi cho các tính toán thống kê có tính cộng tính."
    )
    add_body(doc, sec_221)

    # 2.2.2 Dữ liệu GPR
    add_heading(doc, "2.2.2. Dữ liệu rủi ro địa chính trị (GPR)", level=2)
    sec_222 = (
        "Biến số độc lập chính của nghiên cứu là chỉ số rủi ro địa chính trị (Geopolitical Risk Index – "
        "GPR) do Caldara và Iacoviello phát triển. Chỉ số này đo lường mức độ rủi ro dựa trên tần suất "
        "xuất hiện của các từ khóa liên quan đến căng thẳng quân sự, khủng bố và xung đột ngoại giao trên "
        "một nhóm báo quốc tế lớn.\n\n"
        "Do dữ liệu GPR gốc được công bố theo tần suất tháng, nghiên cứu áp dụng phương pháp gán giá trị "
        "bước (step function) để đồng nhất với dữ liệu ngày của tài sản. Để nắm bắt tính động của thông "
        "tin, biến cú sốc GPR được xây dựng dựa trên sai phân bậc nhất:\n\n"
        "ΔGPR_t = GPR_t - GPR_{t-1}\n\n"
        "Biến ΔGPR_t phản ánh sự thay đổi ròng của mức độ rủi ro, phù hợp với cách tiếp cận “tác động của "
        "thông tin mới” trong các mô hình định giá tài sản."
    )
    add_body(doc, sec_222)

    # 2.3 Wavelet QQR
    add_heading(doc, "2.3. Phương pháp Wavelet Quantile-on-Quantile Regression (Wavelet QQR)", level=1)

    sec_23_intro = (
        "Wavelet QQR là phương pháp trọng tâm để giải mã cấu trúc mối liên hệ phi tuyến và đa thang đo thời "
        "gian giữa GPR và lợi suất BTC, GOLD, OIL. Phương pháp này tích hợp hai kỹ thuật: Biến đổi Wavelet "
        "Rời rạc Chồng lấp Tối đa (MODWT) và Hồi quy Lượng tử trên Lượng tử (QQR)."
    )
    add_body(doc, sec_23_intro)

    # 2.3.1 MODWT
    add_heading(doc, "2.3.1. Phân rã đa độ phân giải với MODWT", level=2)
    sec_231 = (
        "Để khắc phục hạn chế của phân tích thuần túy trong miền thời gian, nghiên cứu áp dụng MODWT để phân "
        "rã chuỗi lợi suất tài sản r_{i,t} và cú sốc rủi ro ΔGPR_t thành các thành phần tần số khác nhau. "
        "Nghiên cứu sử dụng bộ lọc wavelet LA8 (Least Asymmetric độ dài 8) – tương ứng với sym4 trong thư viện "
        "PyWavelets – do khả năng giảm thiểu dịch chuyển pha và cho dạng sóng trơn tru hơn so với các bộ lọc "
        "cơ bản.\n\n"
        "Dữ liệu được phân tách thành J = 9 thang đo (scales), ký hiệu là D_j (j = 1,…,9), tương ứng với các "
        "khung thời gian đầu tư:\n\n"
        "- Ngắn hạn (D_1–D_2): Tương ứng chu kỳ từ khoảng 2–4 ngày đến 4–8 ngày.\n"
        "- Trung hạn (D_3–D_5): Tương ứng chu kỳ từ khoảng 8–32 ngày.\n"
        "- Dài hạn (D_6–D_9): Tương ứng chu kỳ từ khoảng 64–512 ngày, phản ánh các xu hướng chậm và yếu tố vĩ mô.\n\n"
        "Ưu điểm của MODWT so với DWT truyền thống là không yêu cầu giảm mẫu (downsampling), do đó bảo toàn được "
        "độ dài chuỗi dữ liệu gốc và cho phép phân tích chi tiết tại mọi thời điểm."
    )
    add_body(doc, sec_231)

    # 2.3.2 QQR
    add_heading(doc, "2.3.2. Mô hình Hồi quy Lượng tử trên Lượng tử (QQR)", level=2)
    sec_232 = (
        "Tại mỗi thang đo wavelet j, nghiên cứu áp dụng mô hình QQR (Sim & Zhou, 2015) để kiểm định mối quan "
        "hệ phụ thuộc giữa phân vị của lợi suất tài sản và phân vị của cú sốc GPR. Ký hiệu r_{i}^{(j)} là thành "
        "phần wavelet của lợi suất tài sản i ở thang j và ΔGPR^{(j)} là thành phần wavelet của cú sốc GPR, mô "
        "hình có dạng tổng quát:\n\n"
        "Q_θ(r_{i}^{(j)} | ΔGPR^{(j)}) = β_0(θ, τ) + β_1(θ, τ) · (ΔGPR^{(j)} - Q_τ(ΔGPR^{(j)}))\n\n"
        "Trong đó θ là phân vị của lợi suất tài sản (đại diện cho trạng thái thị trường: giảm mạnh, trung tính, "
        "tăng mạnh) và τ là phân vị của cú sốc GPR (đại diện cho mức độ rủi ro: thấp, trung bình, cao). Hệ số "
        "β_1(θ, τ) đo lường tác động biên của GPR lên lợi suất tài sản tại một cặp phân vị cụ thể. Tập hợp các "
        "hệ số này tạo thành một ma trận β(θ, τ), được trực quan hóa dưới dạng bản đồ nhiệt để phân tích cấu "
        "trúc phụ thuộc phi tuyến ở nhiều trạng thái thị trường và nhiều thang thời gian."
    )
    add_body(doc, sec_232)

    # 2.4 Event Study
    add_heading(doc, "2.4. Phương pháp Nghiên cứu Sự kiện (Event Study)", level=1)

    sec_24_intro = (
        "Nghiên cứu sự kiện được sử dụng như một công cụ bổ sung cho Wavelet QQR nhằm phân loại và gắn nhãn "
        "các chế độ phản ứng giá quanh các cú sốc địa chính trị cụ thể. Thay vì chỉ kiểm định xem sự kiện có "
        "tác động hay không, khóa luận tập trung vào hình dạng và dấu của phản ứng giá BTC, GOLD và OIL."
    )
    add_body(doc, sec_24_intro)

    # 2.4.1 Windows and expected returns
    add_heading(doc, "2.4.1. Xác định cửa sổ và mô hình lợi suất kỳ vọng", level=2)
    sec_241 = (
        "Đối với mỗi sự kiện k xảy ra tại ngày t_0, nghiên cứu thiết lập hai khoảng thời gian:\n\n"
        "- Cửa sổ ước lượng (Estimation Window): từ t_0 − 120 đến t_0 − 20 ngày, dùng để ước tính hành vi "
        "“bình thường” của tài sản khi chưa có tin tức.\n"
        "- Cửa sổ sự kiện (Event Window): từ t_0 − 10 đến t_0 + 10 ngày, đủ rộng để nắm bắt cả hiệu ứng rò rỉ "
        "thông tin trước sự kiện và phản ứng điều chỉnh sau khi sự kiện xảy ra.\n\n"
        "Nghiên cứu sử dụng mô hình điều chỉnh trung bình (Mean-adjusted Model) để ước lượng lợi suất kỳ vọng:\n\n"
        "E[r_{i,t}] = (1/L) · Σ_{t ∈ estimation} r_{i,t}\n\n"
        "Lựa chọn này giúp tránh phụ thuộc vào mô hình thị trường phức tạp, phù hợp với đặc thù các tài sản như Bitcoin."
    )
    add_body(doc, sec_241)

    # 2.4.2 AR & CAR
    add_heading(doc, "2.4.2. Tính toán Lợi suất Bất thường (AR) và CAR", level=2)
    sec_242 = (
        "Tác động của sự kiện được đo lường thông qua hai đại lượng cơ bản:\n\n"
        "- Lợi suất bất thường (Abnormal Return – AR): AR_{i,t} = r_{i,t} − E[r_{i,t}].\n"
        "- Lợi suất bất thường tích lũy (Cumulative Abnormal Return – CAR):\n"
        "  CAR_i(t_1, t_2) = Σ_{t=t_1}^{t_2} AR_{i,t}.\n\n"
        "Trong nghiên cứu, CAR được tính cho một số cửa sổ điển hình như (−1, +1), (−3, +3) và (−5, +5) để "
        "so sánh phản ứng rất ngắn hạn và ngắn hạn mở rộng, đồng thời vẽ quỹ đạo CAR trên toàn bộ khoảng "
        "(−10, +10) ngày."
    )
    add_body(doc, sec_242)

    # 2.4.3 Regimes
    add_heading(doc, "2.4.3. Phân loại chế độ phản ứng giá (response regimes)", level=2)
    sec_243 = (
        "Điểm nhấn của khối phân tích sự kiện là sử dụng CAR của BTC và GOLD để phân loại 61 sự kiện vào "
        "các nhóm chế độ phản ứng giá khác nhau. Dựa trên dấu của CAR, mỗi sự kiện được gán vào một trong "
        "bốn nhóm:\n\n"
        "- Regime 1 (BTC↑ – GOLD↑): CAR_{BTC} > 0 và CAR_{GOLD} > 0.\n"
        "- Regime 2 (BTC↓ – GOLD↑): CAR_{BTC} < 0 và CAR_{GOLD} > 0.\n"
        "- Regime 3 (BTC↑ – GOLD↓): CAR_{BTC} > 0 và CAR_{GOLD} < 0.\n"
        "- Regime 4 (BTC↓ – GOLD↓): CAR_{BTC} < 0 và CAR_{GOLD} < 0.\n\n"
        "Việc phân loại này, kết hợp với thông tin về vùng địa lý và loại sự kiện, cho phép liên kết trực tiếp "
        "các kết quả định lượng với bối cảnh địa chính trị cụ thể. Đồng thời, các nhóm regime này được sử dụng "
        "làm “nhãn” để đối chiếu với các vùng phụ thuộc mạnh/yếu trên bản đồ Wavelet QQR."
    )
    add_body(doc, sec_243)

    # 2.5 Visualization
    add_heading(doc, "2.5. Hệ thống trực quan hóa kết quả", level=1)
    sec_25 = (
        "Để kết nối các kết quả định lượng phức tạp với diễn giải kinh tế, khóa luận sử dụng hệ thống trực "
        "quan hóa đa chiều, bao gồm:\n\n"
        "- Biểu đồ Wavelet QQR (heatmap hoặc bề mặt 3D): thể hiện ma trận β(θ, τ) cho từng thang thời gian, "
        "giúp nhận diện các vùng phụ thuộc mạnh hoặc đảo chiều giữa GPR và lợi suất tài sản.\n"
        "- Bản đồ chế độ (Regime Map): biểu đồ phân tán giữa CAR_{BTC} và CAR_{GOLD}, trong đó mỗi điểm là "
        "một sự kiện, màu sắc thể hiện loại sự kiện và ký hiệu thể hiện vùng địa lý.\n"
        "- Bản đồ nhiệt địa chính trị (Geopolitical Heatmap): ma trận màu thể hiện CAR trung bình của từng "
        "tài sản theo từng khu vực, làm nổi bật các vùng nhạy cảm với rủi ro địa chính trị.\n"
        "- Quỹ đạo CAAR (CAAR trajectory): đường diễn biến CAR trung bình quanh ngày sự kiện cho từng nhóm "
        "sự kiện, giúp so sánh tốc độ và hướng phản ứng giá giữa các bối cảnh khác nhau."
    )
    add_body(doc, sec_25)

    # 2.6 Summary
    add_heading(doc, "2.6. Tóm tắt chương", level=1)
    sec_26 = (
        "Chương 2 đã thiết lập khung phương pháp luận cho khóa luận, kết hợp sức mạnh phân tích cấu trúc "
        "sâu của Wavelet QQR với khả năng phân loại và minh họa cụ thể của Nghiên cứu Sự kiện. Cách tiếp "
        "cận đa phương pháp này cho phép khóa luận không chỉ đo lường được mức độ tác động của rủi ro địa "
        "chính trị lên hành động giá của BTC, GOLD và OIL, mà còn giải mã được cơ chế và bối cảnh phản ứng "
        "giữa các vùng địa lý và loại sự kiện. Đây là nền tảng quan trọng để trình bày và thảo luận kết quả "
        "ở các chương tiếp theo."
    )
    add_body(doc, sec_26)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_chapter_two()


