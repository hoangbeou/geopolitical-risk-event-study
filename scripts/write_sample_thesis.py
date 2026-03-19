"""Create a sample thesis draft directly in the FTU template format."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt


def configure_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(13)
    fmt = normal_style.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Cm(1)

    chapter_style = doc.styles.add_style("ChapterTitle", 1)
    chapter_style.font.name = "Times New Roman"
    chapter_style.font.size = Pt(16)
    chapter_style.font.bold = True
    chapter_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    chapter_style.paragraph_format.first_line_indent = Pt(0)

    prelude_style = doc.styles.add_style("PreludeHeading", 1)
    prelude_style.font.name = "Times New Roman"
    prelude_style.font.size = Pt(14)
    prelude_style.font.bold = True
    prelude_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    prelude_style.paragraph_format.first_line_indent = Pt(0)


def add_centered(doc: Document, text: str, style_name: str | None = None) -> None:
    para = doc.add_paragraph(text)
    if style_name:
        para.style = style_name
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.paragraph_format.first_line_indent = Pt(0)


def write_cover(doc: Document) -> None:
    add_centered(doc, "TRƯỜNG ĐẠI HỌC NGOẠI THƯƠNG", "PreludeHeading")
    add_centered(doc, "KHOA KINH TẾ QUỐC TẾ", "PreludeHeading")
    doc.add_paragraph("")
    add_centered(doc, "KHÓA LUẬN TỐT NGHIỆP", "PreludeHeading")
    add_centered(
        doc,
        "TÁC ĐỘNG CỦA RỦI RO ĐỊA CHÍNH TRỊ ĐẾN BTC, VÀNG VÀ DẦU THÔ",
        "PreludeHeading",
    )
    add_centered(doc, "Chuyên ngành: Kinh tế quốc tế")
    add_centered(doc, "Giảng viên hướng dẫn: TS. Nguyễn An Nhiên")
    add_centered(doc, "Sinh viên thực hiện: Lê Minh Anh")
    add_centered(doc, "Mã sinh viên: 2012410010")
    add_centered(doc, f"Hà Nội, {date.today().year}")
    doc.add_page_break()


def write_front_matter(doc: Document) -> None:
    doc.add_paragraph("LỜI CẢM ƠN", style="ChapterTitle")
    doc.add_paragraph(
        "Khóa luận được hoàn thành nhờ sự hướng dẫn tận tình của TS. Nguyễn An Nhiên "
        "và sự hỗ trợ dữ liệu từ Trung tâm Nghiên cứu Kinh tế số. Tác giả xin trân trọng "
        "cảm ơn thầy cô, gia đình và các đồng môn đã đồng hành trong suốt quá trình nghiên cứu."
    )

    doc.add_paragraph("MỤC LỤC", style="ChapterTitle")
    doc.add_paragraph(
        "Mục lục sẽ được cập nhật tự động sau khi hoàn thiện bản thảo cuối cùng."
    )
    doc.add_paragraph("DANH MỤC CÁC CHỮ VIẾT TẮT", style="ChapterTitle")
    doc.add_paragraph(
        "BTC: Bitcoin; GPR: Geopolitical Risk Index; CAR: Cumulative Abnormal Return; "
        "QQR: Quantile-on-Quantile Regression; LA8: Least Asymmetric wavelet."
    )
    doc.add_paragraph("DANH MỤC CÁC BẢNG BIỂU", style="ChapterTitle")
    doc.add_paragraph(
        "Bảng 1.1 Tóm tắt các nghiên cứu về tài sản trú ẩn; Bảng 3.2 CAR trung bình theo khu vực; "
        "Bảng 4.1 So sánh cơ chế phản ứng tại châu Á và châu Âu."
    )
    doc.add_paragraph("DANH MỤC CÁC HÌNH VẼ", style="ChapterTitle")
    doc.add_paragraph(
        "Hình 3.1 Regime Map BTC-GOLD; Hình 3.2 Geopolitical Heatmap; "
        "Hình 4.1 Regional Pattern Radar."
    )

    doc.add_paragraph("TÓM TẮT KHÓA LUẬN", style="ChapterTitle")
    doc.add_paragraph(
        "Khóa luận đánh giá 61 sự kiện địa chính trị giai đoạn 2013–2025 để đo lường phản ứng "
        "của Bitcoin (BTC), vàng (GOLD) và dầu thô (OIL). Phương pháp nghiên cứu gồm Event Study "
        "với cửa sổ ±10 ngày và Wavelet Quantile-on-Quantile Regression (QQR) sử dụng bộ lọc LA8 "
        "với 9 thang thời gian. Kết quả cho thấy BTC có 4 chế độ phản ứng khác nhau với GOLD, "
        "trong đó cơ chế vùng đóng vai trò quyết định: châu Á thường chứng kiến BTC↑GOLD↑ "
        "do kỳ vọng nới lỏng kiểm soát vốn, trong khi châu Âu xuất hiện BTC↓GOLD↓ vì cú sốc "
        "lan rộng sang hệ thống tài chính. Dầu thô gần như không chịu tác động đáng kể từ GPR. "
        "Kết quả bổ sung bằng QQR khẳng định phụ thuộc phi tuyến và khác biệt giữa các phân vị "
        "cực đoan, qua đó đề xuất chiến lược phòng ngừa rủi ro theo vùng và loại sự kiện."
    )


def write_intro(doc: Document) -> None:
    doc.add_paragraph("PHẦN MỞ ĐẦU", style="ChapterTitle")
    doc.add_paragraph(
        "1. Tính cấp thiết của nghiên cứu: Biến động địa chính trị gia tăng mạnh từ 2020 "
        "đến nay khiến nhà đầu tư tìm kiếm tài sản trú ẩn mới. Việt Nam thiếu nghiên cứu cập nhật "
        "về cách BTC, vàng và dầu phản ứng với các cú sốc này, đặc biệt trong bối cảnh thị trường "
        "crypto nội địa phát triển nhanh."
    )
    doc.add_paragraph(
        "2. Mục tiêu: (i) Phân loại và đo lường phản ứng của ba tài sản trước từng loại sự kiện "
        "địa chính trị; (ii) Phân tích khác biệt khu vực và cơ chế truyền dẫn; (iii) Kiểm định "
        "mối liên hệ định lượng giữa GPR và lợi suất thông qua Wavelet QQR."
    )
    doc.add_paragraph(
        "3. Đối tượng và phạm vi: 61 sự kiện toàn cầu ảnh hưởng trực tiếp đến an ninh, thương mại "
        "hoặc quan hệ quốc tế giai đoạn 2013–11/2025. Dữ liệu giá BTC, vàng COMEX và dầu Brent "
        "được thu thập theo ngày; chỉ số GPR, GPRD_ACT, GPRD_THREAT lấy từ cơ sở dữ liệu Caldara-Iacoviello."
    )
    doc.add_paragraph(
        "4. Phương pháp và dữ liệu: Event Study (mean-adjusted model) để tính CAR ±10 ngày; "
        "Wavelet MODWT với bộ lọc LA8 và 9 thang để tách tín hiệu; QQR chiếu phân vị GPR vào "
        "phân vị lợi suất nhằm nhận diện phụ thuộc phi tuyến. Các thống kê mô tả và kiểm định "
        "robustness được thực hiện bằng Python."
    )
    doc.add_paragraph(
        "5. Cấu trúc khóa luận: Chương 1 tổng quan lý thuyết; Chương 2 mô tả dữ liệu – phương pháp; "
        "Chương 3 trình bày kết quả event study và QQR; Chương 4 thảo luận cơ chế vùng – loại sự kiện; "
        "Chương 5 đề xuất giải pháp và hàm ý chính sách."
    )


def write_chapter(
    doc: Document, code: str, title: str, sections: list[tuple[str, str]]
) -> None:
    doc.add_page_break()
    doc.add_paragraph(code, style="ChapterTitle")
    doc.add_paragraph(title, style="ChapterTitle")
    for heading, body in sections:
        doc.add_paragraph(heading)
        for paragraph in body.split("\n"):
            doc.add_paragraph(paragraph.strip())


def build_document() -> Document:
    doc = Document()
    configure_layout(doc)
    write_cover(doc)
    write_front_matter(doc)
    write_intro(doc)

    chapter1_sections = [
        (
            "1.1 Tổng quan nghiên cứu",
            "Các nghiên cứu về tài sản trú ẩn (Baur & Lucey, 2010; Dyhrberg, 2016) cho thấy vàng "
            "và trái phiếu Mỹ là nơi trú ẩn truyền thống, trong khi Bitcoin có tính chất lai. "
            "Trong bối cảnh địa chính trị, Al-Thaqeb et al. (2022) chỉ ra rằng cú sốc chiến tranh "
            "làm tăng biến động BTC nhưng không thống nhất về hướng phản ứng. Phần lớn tài liệu "
            "tập trung vào COVID-19, ít chú trọng giai đoạn 2022–2025 với xung đột Nga–Ukraine, "
            "Israel–Hamas hay căng thẳng Đài Loan.\n"
            "Khoảng trống lớn nằm ở phân tích phối hợp BTC–GOLD–OIL theo vùng và loại sự kiện. "
            "Các nghiên cứu wavelet trước đây (Yarovaya et al., 2022) chưa xem xét GPRD_ACT/THREAT "
            "và chưa phân rã phản ứng tức thời của từng tài sản.",
        ),
        (
            "1.2 Cơ sở lý thuyết",
            "Khung lý thuyết dựa trên giả thuyết tài sản trú ẩn và lý thuyết kênh truyền dẫn rủi ro. "
            "GPR tác động tới dòng vốn qua hai kênh: (i) kỳ vọng phòng thủ (flight-to-quality) khiến "
            "vàng tăng; (ii) né tránh kiểm soát vốn, thúc đẩy tài sản số. Lý thuyết thông tin bất cân "
            "xứng giải thích tại sao các sự kiện đe dọa (THREAT) làm tăng tài sản rủi ro nếu nhà đầu "
            "tư kỳ vọng sự kiện không hiện thực hóa.\n"
            "Khung nghiên cứu áp dụng mô hình CAR để đo phản ứng tức thì và Wavelet QQR để kiểm tra "
            "phụ thuộc phi tuyến theo phân vị, phù hợp với bản chất dữ liệu tài chính nhiều cực trị.",
        ),
        (
            "1.3 Khoảng trống nghiên cứu",
            "Chưa có nghiên cứu nào kết hợp Event Study, GPRD_ACT/THREAT và Wavelet QQR để lý giải "
            "đồng thời hai hiện tượng: (i) BTC đôi khi di chuyển cùng chiều vàng, đôi khi ngược chiều; "
            "(ii) sự khác biệt rõ ràng giữa châu Á và châu Âu. Khóa luận nhằm lấp khoảng trống này "
            "và cung cấp thêm chứng cứ cho thị trường mới nổi như Việt Nam.",
        ),
    ]

    chapter2_sections = [
        (
            "2.1 Thiết kế nghiên cứu",
            "Nghiên cứu sử dụng hai giai đoạn: Event Study để xác định phản ứng tức thì và Wavelet QQR "
            "để kiểm tra mối liên hệ dài – ngắn hạn. Cửa sổ sự kiện được cố định ±10 ngày quanh ngày t=0 "
            "của 61 sự kiện. Các sự kiện được phân loại thành War, Political, Mixed và gắn vùng phát sinh "
            "(Asia, Europe, Middle East, Americas, Global).",
        ),
        (
            "2.2 Dữ liệu",
            "Giá BTC (CoinMarketCap), giá vàng (London Fix, qui đổi COMEX) và giá dầu Brent được log "
            "hóa rồi tính lợi suất theo ngày. Chỉ số GPR tổng hợp và hai thành phần ACT/THREAT "
            "được lấy từ dữ liệu Caldara–Iacoviello (bản cập nhật 09/2025). Ngoài ra, khóa luận xây dựng "
            "tập sự kiện nội bộ dựa trên Reuters, AP, NATO briefings nhằm bảo đảm đủ 61 sự kiện.",
        ),
        (
            "2.3 Phương pháp",
            "Bước 1: Ước lượng lợi suất kỳ vọng bằng mô hình trung bình lịch sử trong cửa sổ 120 ngày trước sự kiện. "
            "Bước 2: Tính CAR cho từng tài sản và phân loại dấu (+/-). Bước 3: Thực hiện MODWT với bộ lọc LA8, "
            "9 thang để thu các thành phần dao động. Bước 4: Áp dụng QQR với lưới phân vị 0.05–0.95, kernel "
            "Gaussian và bandwidth Silverman. Các kiểm định độ tin cậy gồm phân tích độ nhạy cửa sổ sự kiện "
            "và kiểm tra Heteroskedasticity-autocorrelation robust (HAR).",
        ),
    ]

    chapter3_sections = [
        (
            "3.1 Phản ứng theo loại sự kiện",
            "BTC tăng trong 32/61 sự kiện (CAR trung bình +16.6%), chủ yếu là đe dọa quân sự ở châu Á và hỗn hợp "
            "thương mại tại châu Mỹ. BTC giảm trong 29 sự kiện, tập trung ở các pha leo thang thực chiến như Nga–Ukraine "
            "2022 hoặc Israel–Hamas 2023. Vàng tăng mạnh trong các sự kiện chiến tranh kéo dài, trong khi dầu chỉ phản ứng "
            "đáng kể với các cú sốc nguồn cung thực tế (ví dụ tấn công cơ sở dầu mỏ Ả Rập).",
        ),
        (
            "3.2 So sánh BTC – GOLD – OIL",
            "Có bốn pattern chính: cả hai cùng tăng (29.5%), cùng giảm (24.6%), BTC↑GOLD↓ (23.0%) và BTC↓GOLD↑ (23.0%). "
            "Pattern BTC↑GOLD↓ thường đi kèm sự kiện mang tính đe dọa mơ hồ tại châu Á, minh họa vai trò của BTC như tài sản "
            "thay thế khi nhà đầu tư kỳ vọng kiểm soát vốn. Ngược lại, pattern BTC↓GOLD↑ xuất hiện khi sự kiện làm nổi bật "
            "chức năng trú ẩn truyền thống của vàng.\n"
            "OIL thể hiện CAR quanh 0 trong phần lớn trường hợp, xác nhận rằng rủi ro địa chính trị chỉ tác động mạnh khi "
            "liên quan trực tiếp đến chuỗi cung ứng năng lượng.",
        ),
        (
            "3.3 Kết quả Wavelet QQR",
            "QQR cho thấy phụ thuộc phi tuyến rõ rệt: tại các phân vị GPR cao (>0.8) và phân vị lợi suất thấp (<0.2), "
            "BTC mang hệ số âm có ý nghĩa, hàm ý cú sốc thực chiến khiến BTC giảm. Tuy nhiên ở phân vị lợi suất cao (>0.8), "
            "các cú sốc đe dọa lại làm hệ số dương mạnh, củng cố luận điểm \"buy the rumor\". Vàng thể hiện hệ số dương "
            "ổn định trên các phân vị, còn dầu vẫn dao động quanh 0. Tính đa thang cho thấy phản ứng của BTC diễn ra nhanh "
            "ở thang 1–2 (2–4 ngày), trong khi vàng hấp thụ thông tin ở thang 3–4 (4–8 ngày).",
        ),
    ]

    chapter4_sections = [
        (
            "4.1 Tổng hợp các pattern",
            "Việc chồng ghép CAR của BTC và GOLD tạo thành Regime Map với bốn góc phần tư rõ ràng. Góc BTC↑GOLD↓ "
            "quy tụ các sự kiện bị thị trường coi là cơ hội thay thế tài sản (ví dụ sắc lệnh siết vốn tại Trung Quốc 2024), "
            "trong khi góc BTC↓GOLD↑ là minh chứng cho flight-to-quality (khủng hoảng Gaza 2023).",
        ),
        (
            "4.2 Cơ chế khu vực",
            "Châu Á: BTC↑GOLD↑ mạnh nhất (BTC +22.1%, GOLD +25.0%). Nguyên nhân chính là nguy cơ kiểm soát vốn "
            "và sự tham gia mạnh của nhà đầu tư cá nhân, khiến cả BTC và vàng được gom mua như bộ đôi phòng vệ. "
            "Các sự kiện tiêu biểu gồm căng thẳng eo biển Đài Loan và lệnh giới nghiêm tại Hong Kong.\n"
            "Châu Âu: BTC↓GOLD↓ phổ biến (8 sự kiện). Khi chiến sự Nga–Ukraine leo thang, rủi ro hệ thống làm BTC giảm "
            "theo tâm lý risk-off còn vàng cũng chịu áp lực chốt lời vì nhu cầu thanh khoản USD. Trung Đông lại thể hiện "
            "BTC↓GOLD↓ với biên độ sâu hơn do nhà đầu tư chuyển sang dầu và USD. Mỹ và châu Mỹ chứng kiến BTC↑GOLD↓ "
            "bởi các sự kiện chính trị thường ảnh hưởng kỳ vọng chính sách tiền tệ hơn là đe dọa vật chất.",
        ),
        (
            "4.3 So sánh với nghiên cứu trước",
            "Kết quả tương đồng với Yarovaya et al. (2022) ở điểm BTC không phải trú ẩn ổn định, nhưng khóa luận bổ sung "
            "chứng cứ rằng yếu tố vùng mới là chìa khóa giải thích sự bất định. Khác với luận văn Greenbond–BTC trước đây, "
            "nghiên cứu này thay đổi bộ lọc LA8 và mở rộng giai đoạn 2024–2025, qua đó hệ số QQR nhỏ hơn trung bình "
            "nhưng vẫn có ý nghĩa ở phân vị cực đoan.",
        ),
    ]

    chapter5_sections = [
        (
            "5.1 Khuyến nghị cho nhà đầu tư",
            "Nhà đầu tư nên theo dõi phân loại sự kiện (Threat vs Action) và vùng xảy ra để điều chỉnh vị thế BTC/GOLD. "
            "Khi sự kiện mang tính đe dọa mơ hồ tại châu Á, chiến lược long BTC kết hợp hedge vàng mang lại CAR dương. "
            "Ngược lại, khi chiến sự ở châu Âu leo thang, cần giảm tỷ trọng BTC và chuyển sang vàng-ngắn hạn hoặc USD.",
        ),
        (
            "5.2 Hàm ý cho nhà hoạch định chính sách",
            "Cơ quan quản lý cần chuẩn bị kịch bản biến động tiền kỹ thuật số khi xảy ra khủng hoảng khu vực. Việc giám sát "
            "dòng vốn vào BTC trong các giai đoạn căng thẳng sẽ giúp ổn định thị trường ngoại hối. Ngoài ra, minh bạch thông tin "
            "về dự trữ vàng quốc gia có thể củng cố niềm tin nhà đầu tư.",
        ),
        (
            "5.3 Hạn chế và hướng nghiên cứu",
            "Khóa luận chưa đánh giá sâu các kênh truyền dẫn vĩ mô (lãi suất, tỷ giá) và chưa kiểm định mô hình nhân quả "
            "đa biến. Hướng nghiên cứu tiếp theo có thể mở rộng sang tài sản trái phiếu xanh hoặc sử dụng dữ liệu tần suất "
            "cao để cải thiện độ chính xác của Event Study.",
        ),
    ]

    write_chapter(doc, "CHƯƠNG 1", "CƠ SỞ LÝ LUẬN VÀ TỔNG QUAN", chapter1_sections)
    write_chapter(doc, "CHƯƠNG 2", "DỮ LIỆU VÀ PHƯƠNG PHÁP", chapter2_sections)
    write_chapter(doc, "CHƯƠNG 3", "KẾT QUẢ NGHIÊN CỨU", chapter3_sections)
    write_chapter(doc, "CHƯƠNG 4", "THẢO LUẬN CƠ CHẾ", chapter4_sections)
    write_chapter(doc, "CHƯƠNG 5", "HÀM Ý VÀ GIẢI PHÁP", chapter5_sections)

    doc.add_page_break()
    doc.add_paragraph("KẾT LUẬN", style="ChapterTitle")
    doc.add_paragraph(
        "Khóa luận chứng minh phản ứng của BTC và vàng trước rủi ro địa chính trị phụ thuộc mạnh vào đặc điểm vùng "
        "và bản chất sự kiện. Phương pháp kết hợp Event Study và Wavelet QQR giúp nắm bắt cả tác động tức thời lẫn "
        "phụ thuộc phi tuyến. Những phát hiện về cơ chế châu Á vs châu Âu mở ra hướng tiếp cận mới cho quản lý danh mục "
        "và hoạch định chính sách tại Việt Nam."
    )

    doc.add_paragraph("DANH MỤC TÀI LIỆU THAM KHẢO", style="ChapterTitle")
    doc.add_paragraph("Baur, D. G., & Lucey, B. M. (2010). Is gold a hedge or a safe haven?")
    doc.add_paragraph(
        "Caldara, D., & Iacoviello, M. (2022). Measuring geopolitical risk."
    )
    doc.add_paragraph(
        "Yarovaya, L., et al. (2022). Geopolitical risk and cryptocurrencies: "
        "Evidence from wavelet analysis."
    )

    doc.add_paragraph("PHỤ LỤC", style="ChapterTitle")
    doc.add_paragraph("Phụ lục A: Danh sách 61 sự kiện; Phụ lục B: Quy trình tính CAR.")
    return doc


def main() -> None:
    doc = build_document()
    target = Path("docs/thesis/templates/KLTN_FORMAT_SAMPLE.docx")
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    print(f"[OK] Sample thesis saved to {target}")


if __name__ == "__main__":
    main()

