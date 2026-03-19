# -*- coding: utf-8 -*-
"""
Generate a compact Word document (~88 pages) by keeping only essential chapters.
"""
from __future__ import annotations

from pathlib import Path
from typing import Iterable, List

from docx import Document

THESIS_DIR = Path("docs/thesis")
FRONT_MATTER_FILE = "KHOA_LUAN_HOAN_CHINH.md"
CHAPTER_FILES = [
    "KHOA_LUAN_CHUONG_1.md",
    "KHOA_LUAN_CHUONG_2.md",
    "KHOA_LUAN_CHUONG_3.md",
    "KHOA_LUAN_CHUONG_4.md",
    "KHOA_LUAN_CHUONG_5.md",
]


def read_lines(path: Path) -> List[str]:
    return path.read_text(encoding="utf-8").splitlines()


def extract_section(path: Path, heading: str) -> List[str]:
    lines = read_lines(path)
    extracted: List[str] = []
    capture = False
    heading_lower = heading.strip().lower()

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("##"):
            title = stripped.lstrip("#").strip(" *").lower()
            if title.startswith(heading_lower):
                capture = True
                continue
            elif capture:
                break
        if capture:
            extracted.append(line)

    return extracted


def add_markdown_lines(doc: Document, lines: Iterable[str]) -> None:
    in_code_block = False
    code_buffer: List[str] = []

    for line in lines:
        stripped = line.rstrip("\n")

        if stripped.strip().startswith("```"):
            if in_code_block:
                doc.add_paragraph("\n".join(code_buffer), style="Intense Quote")
                code_buffer = []
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_buffer.append(stripped)
            continue

        if not stripped.strip():
            doc.add_paragraph("")
            continue

        if stripped.lstrip().startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("#").strip()
            doc.add_heading(text, level=min(level, 4))
            continue

        if stripped.strip().startswith(("- ", "* ")):
            doc.add_paragraph(stripped.strip()[2:], style="List Bullet")
            continue

        if stripped.strip().startswith(">"):
            doc.add_paragraph(stripped.strip()[1:].strip(), style="Intense Quote")
            continue

        doc.add_paragraph(stripped)

    if code_buffer:
        doc.add_paragraph("\n".join(code_buffer), style="Intense Quote")


def add_markdown_file(doc: Document, path: Path, page_break_before: bool = False) -> None:
    if page_break_before:
        doc.add_page_break()
    add_markdown_lines(doc, read_lines(path))


def build_compact_doc(output: Path) -> None:
    doc = Document()
    doc.add_heading(
        "PHÂN TÍCH ẢNH HƯỞNG CỦA RỦI RO ĐỊA CHÍNH TRỊ ĐẾN BTC, GOLD, OIL",
        level=0,
    )
    doc.add_paragraph("Khóa luận tốt nghiệp – phiên bản rút gọn trong 88 trang.")
    doc.add_paragraph(
        "Tài liệu này giữ lại các nội dung cốt lõi (Tóm tắt + Chương 1-5) và "
        "loại bỏ phần hướng dẫn/dàn ý để giảm số trang.",
        style="Intense Quote",
    )

    front_path = THESIS_DIR / FRONT_MATTER_FILE
    for heading in ["TÓM TẮT", "TÓM TẮT NỘI DUNG", "KẾT QUẢ CHÍNH", "Ý NGHĨA"]:
        section_lines = extract_section(front_path, heading)
        if section_lines:
            doc.add_heading(heading.title(), level=1)
            add_markdown_lines(doc, section_lines)

    for idx, chapter in enumerate(CHAPTER_FILES):
        add_markdown_file(doc, THESIS_DIR / chapter, page_break_before=True)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Saved compact thesis to {output}")


if __name__ == "__main__":
    out_file = THESIS_DIR / "KHOA_LUAN_COMPACT.docx"
    build_compact_doc(out_file)

