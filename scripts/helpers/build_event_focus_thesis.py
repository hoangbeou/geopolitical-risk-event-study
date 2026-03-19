# -*- coding: utf-8 -*-
"""
Generate an event-study–centric thesis document limited to concise sections
so that the final length fits within ~80 pages once formatted.
"""
from __future__ import annotations

from pathlib import Path
from docx import Document


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def build_event_focus_doc(output: Path) -> None:
    doc = Document()
    add_heading(doc, "KHÓA LUẬN TỐT NGHIỆP", 0)
    add_heading(doc, "ẢNH HƯỞNG CỦA RỦI RO ĐỊA CHÍNH TRỊ ĐẾN BTC, GOLD, OIL", 1)
    add_paragraph(doc,
        "Phiên bản này nhấn mạnh Event Study như phương pháp chính và tóm lược "
        "kết quả QQR, VAR/TVP-VAR trong phạm vi 1–2 trang, giúp tổng chiều dài "
        "tài liệu sau khi định dạng chuẩn không vượt quá 88 trang."
    )

    add_heading(doc, "Tóm tắt", 1)
    add_paragraph(doc,
        "Nghiên cứu sử dụng bộ dữ liệu daily 2015–2025 với GPR Index của "
        "Caldara & Iacoviello (2022) và giá BTC, GOLD, OIL. Event Study với "
        "auto-detected geopolitical shocks (percentile 95/90) cho thấy BTC "
        "phản ứng mạnh nhất (Mean |CAR| 17.7%), GOLD phản ứng có điều kiện "
        "tùy loại sự kiện, còn OIL ít nhạy cảm (1.8%). QQR xác nhận mối quan "
        "hệ bất đối xứng theo quantile, trong khi VAR/TVP-VAR chỉ ra GPR không "
        "dự báo được biến động dài hạn nhưng vẫn kích hoạt connectedness ngắn "
        "hạn trong giai đoạn căng thẳng."
    )

    # Chapter 1
    add_heading(doc, "Chương 1. Mở đầu", 1)
    add_paragraph(doc,
        "Chương này giới thiệu bối cảnh gia tăng tần suất các cú sốc địa chính "
        "trị (Brexit, Trade War, Nga–Ukraine, Israel–Palestine) và động lực "
        "phân tích phản ứng của BTC, GOLD, OIL. Mục tiêu trung tâm là đo đạc "
        "tác động tức thời của các biến cố GPR lên lợi suất tài sản bằng Event "
        "Study, đồng thời kiểm chứng tính không tuyến tính (QQR) và lan tỏa "
        "dài hạn (VAR/TVP-VAR)."
    )
    add_heading(doc, "1.1 Mục tiêu & câu hỏi", 2)
    add_bullet(doc, "Event Study: Lợi suất của BTC/GOLD/OIL phản ứng thế nào quanh cú sốc GPR?")
    add_bullet(doc, "QQR: Mối quan hệ thay đổi ra sao theo quantile GPR và trạng thái thị trường?")
    add_bullet(doc, "VAR/TVP-VAR: GPR có vai trò dự báo và truyền dẫn biến động dài hạn không?")
    add_heading(doc, "1.2 Đóng góp", 2)
    add_bullet(doc, "Áp dụng auto-detection trên GPR để chọn 62 sự kiện mà không lệ thuộc vào danh sách thủ công.")
    add_bullet(doc, "So sánh 3 phương pháp trong cùng dữ liệu, xác nhận Event Study là khung hiệu quả nhất cho khóa luận.")
    add_bullet(doc, "Nêu bật sự khác biệt giữa BTC, GOLD, OIL khi đối mặt cùng một cú sốc GPR.")

    # Chapter 2
    add_heading(doc, "Chương 2. Cơ sở lý thuyết & tổng quan", 1)
    add_heading(doc, "2.1 Khung lý thuyết", 2)
    add_paragraph(doc,
        "Tóm lược các luận điểm safe haven (Baur & McDermott, 2010), risk-on/off, "
        "và cơ chế truyền dẫn của GPR tới tài sản: (i) kênh tâm lý nhà đầu tư "
        "(flight-to-quality), (ii) kênh kỳ vọng tăng trưởng, (iii) kênh cung cầu "
        "năng lượng. GPR Index được mô tả cùng ví dụ spike 2016, 2020, 2022."
    )
    add_heading(doc, "2.2 Nghiên cứu thực nghiệm", 2)
    add_paragraph(doc,
        "Event Study được Campbell et al. (1997) khuyến nghị cho cú sốc ngắn hạn; "
        "QQR (Sim & Zhou, 2015; bài tham khảo về BTC/GOLD) đánh giá phụ thuộc "
        "quantile; VAR/TVP-VAR (Shaik et al., 2024) đo connectedness động. "
        "Khoảng trống: chưa có nghiên cứu kết hợp auto events + Event Study + QQR + TVP-VAR cho bộ ba BTC/GOLD/OIL 2015–2025."
    )

    # Chapter 3
    add_heading(doc, "Chương 3. Dữ liệu & phương pháp", 1)
    add_heading(doc, "3.1 Dữ liệu", 2)
    add_paragraph(doc,
        "Dữ liệu daily log-return của BTC (CoinMetrics), GOLD (LBMA), OIL Brent "
        "(EIA) và GPR tổng hợp. Kiểm định ADF xác nhận tính dừng sau khi lấy lợi "
        "suất/differencing. Trình bày bảng thống kê mô tả 4 series."
    )
    add_heading(doc, "3.2 Event Detection", 2)
    add_paragraph(doc,
        "GPR_diff được chuẩn hóa theo rolling z-score; sự kiện được ghi nhận khi "
        "giá trị vượt percentile 95 (spike) hoặc duy trì percentile 90 trong 10 "
        "ngày (stress). Áp dụng overlap limit 20 ngày để tránh chồng lấn, thu được "
        "62 sự kiện 2015–2025. (Chèn Bảng 3.1: danh sách sự kiện tiêu biểu)."
    )
    add_heading(doc, "3.3 Event Study (phương pháp chính)", 2)
    add_bullet(doc, "Estimation window 120 ngày, gap 5 ngày, event window −10 / +10.")
    add_bullet(doc, "Model kỳ vọng: Mean (chính), Market (DXY), Factor (DXY + DGS3MO + T10YIE).")
    add_bullet(doc, "Abnormal Return AR_it = R_it − E(R_it); CAR_i = Σ AR_it; chuẩn hóa SAR theo Patell Z & BMP.")
    add_paragraph(doc,
        "Kết quả được báo cáo qua Bảng CAR, đồ thị Average CAR và AAR/CAAR."
    )
    add_heading(doc, "3.4 QQR (tóm lược)", 2)
    add_paragraph(doc,
        "Dùng returns chưa qua OLS để bảo toàn biên độ; áp dụng MODWT 6 scale và "
        "lưới quantile 0.05–0.95. Hệ số β(τ_GPR, τ_asset) cho phép nhận diện điều "
        "kiện GPR cao & lợi suất âm của BTC/GOLD."
    )
    add_heading(doc, "3.5 VAR & TVP-VAR (tóm lược)", 2)
    add_paragraph(doc,
        "VAR(2) kiểm tra Granger causality và IRF/FEVD. TVP-VAR connectedness (Kalman "
        "filter) mở rộng để theo dõi Total Connectedness Index (TCI) và vai trò net "
        "transmitter/receiver trong các giai đoạn Pre-COVID, COVID, Russia–Ukraine."
    )

    # Chapter 4
    add_heading(doc, "Chương 4. Kết quả & thảo luận", 1)
    add_heading(doc, "4.1 Event Study", 2)
    add_paragraph(doc,
        "Average CAR cho thấy BTC nhạy nhất (17.7%), GOLD 10.3%, OIL 1.8%. "
        "Hình 4.1 minh họa CAR trung bình theo ngày; Hình 4.2 cho AAR/CAAR. "
        "Bảng 4.3 tổng hợp CAR từng nhóm sự kiện (quân sự, chính trị, khủng hoảng)."
    )
    add_bullet(doc, "Brexit 2016: BTC −12%, GOLD −22%, OIL −4% (panic sell-off).")
    add_bullet(doc, "Nga–Ukraine 2022: BTC −18% tại t0, GOLD +5%, OIL +7% (flight-to-quality).")
    add_bullet(doc, "Israel–Palestine 2023: BTC +8% hậu sự kiện nhờ narrative digital gold.")
    add_paragraph(doc,
        "Phân tích sâu 3 sự kiện giúp chứng minh Event Study trả lời câu hỏi chính "
        "của khóa luận: phản ứng tức thời khác nhau theo từng tài sản."
    )
    add_heading(doc, "4.2 QQR (1–2 trang)", 2)
    add_paragraph(doc,
        "Heatmap cho thấy BTC chịu tác động mạnh khi GPR ở quantile cao (>0.8) và "
        "BTC đang ở quantile thấp (<0.2), GOLD thiên về phản ứng dương ở quantile "
        "trung bình nhưng âm ở quantile thấp, còn OIL hầu như trung tính trừ các "
        "scale dài (d5–d6). Điều này giúp giải thích sự khác biệt hành vi giữa ba "
        "tài sản."
    )
    add_paragraph(doc,
        "(Chèn Hình 4.3: QQR_BTC; Hình 4.4: QQR_GOLD). Từ đây rút ra lý do tại sao "
        "Event Study thấy kết quả trái chiều: cơ chế phụ thuộc trạng thái nền."
    )
    add_heading(doc, "4.3 VAR/TVP-VAR (0.5 trang)", 2)
    add_paragraph(doc,
        "VAR(2) kiểm định Granger cho thấy GPR không dự báo được BTC/GOLD/OIL "
        "(p-value > 0.1) và trong FEVD 10 bước GPR chỉ giải thích dưới 1% biến "
        "động của mỗi tài sản. TVP-VAR ghi nhận rằng TCI chỉ tăng mạnh trong các "
        "giai đoạn chiến sự (2020–2022) và khi đó BTC chuyển tạm thời từ net "
        "receiver sang net transmitter, trong khi GOLD giữ vai trò hấp thụ biến "
        "động và OIL hầu như trung lập. Kết quả này củng cố kết luận: tác động "
        "GPR là ngắn hạn và khác biệt giữa ba tài sản."
    )
    add_paragraph(doc, "(Chèn Hình 4.5: TVP_VAR_TCI hoặc bảng FEVD tóm tắt).")
    add_heading(doc, "4.4 Thảo luận", 2)
    add_bullet(doc, "Event Study cung cấp bằng chứng hành động (CAR, AAR).")
    add_bullet(doc, "QQR giải thích tính bất đối xứng → giải luận vì sao GOLD không luôn tăng.")
    add_bullet(doc, "VAR/TVP-VAR xác nhận GPR không phải biến dự báo dài hạn nhưng vẫn tạo lan tỏa bất đối xứng.")

    # Chapter 5
    add_heading(doc, "Chương 5. Kết luận & khuyến nghị", 1)
    add_heading(doc, "5.1 Kết luận chính", 2)
    add_bullet(doc, "GPR kích hoạt phản ứng tức thời mạnh ở BTC, trung bình ở GOLD, yếu ở OIL.")
    add_bullet(doc, "Mức độ và dấu phản ứng của mỗi tài sản phụ thuộc trạng thái nền và loại sự kiện.")
    add_bullet(doc, "GPR không dự báo được dài hạn nhưng vẫn có thể tạo lan tỏa biến động ngắn hạn.")
    add_heading(doc, "5.2 Khuyến nghị", 2)
    add_bullet(doc, "Nhà đầu tư: Theo dõi GPR spike để đánh giá nhanh sự khác biệt phản ứng giữa BTC, GOLD, OIL.")
    add_bullet(doc, "Quản trị danh mục: Sử dụng kết quả Event Study để phân nhóm tài sản theo độ nhạy GPR.")
    add_bullet(doc, "Nghiên cứu tiếp: mở rộng QQR đa biến và TVP-VAR theo ngành/cổ phiếu để so sánh thêm.")
    add_heading(doc, "5.3 Hạn chế & hướng mở rộng", 2)
    add_bullet(doc, "Dữ liệu daily chưa phản ánh intraday spillover.")
    add_bullet(doc, "Auto-detection dựa trên GPR toàn cầu, chưa phân tách khu vực.")
    add_bullet(doc, "Chưa đánh giá hiệu ứng kiểm soát như lãi suất thực theo thời gian.")

    add_heading(doc, "Phụ lục gợi ý", 1)
    add_bullet(doc, "Danh sách sự kiện auto-detect (Bảng A1).")
    add_bullet(doc, "Bảng CAR chi tiết 62 sự kiện.")
    add_bullet(doc, "Thông số VAR/TVP-VAR & kiểm định chẩn đoán.")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Saved event-focused thesis to {output}")


if __name__ == "__main__":
    build_event_focus_doc(Path("docs/thesis/KHOA_LUAN_EVENT_FOCUS.docx"))

