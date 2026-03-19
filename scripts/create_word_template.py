"""Generate a Word template that follows the FTU KLTN formatting guide."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt


def build_template(output_file: Path) -> None:
    """Create the thesis template document."""

    doc = Document()

    # Page layout: margins per FTU requirement
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2)

    # Base font/paragraph settings
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(13)
    p_format = normal_style.paragraph_format
    p_format.line_spacing = 1.5
    p_format.space_after = Pt(0)
    p_format.space_before = Pt(0)
    p_format.first_line_indent = Cm(1)

    # Custom styles for main titles and helper notes
    chapter_style = doc.styles.add_style("ChapterTitle", 1)
    chapter_style.font.name = "Times New Roman"
    chapter_style.font.size = Pt(16)
    chapter_style.font.bold = True
    chapter_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    chapter_style.paragraph_format.first_line_indent = Pt(0)

    prelude_style = doc.styles.add_style("PreludeHeading", 1)
    prelude_style.font.name = "Times New Roman"
    prelude_style.font.size = Pt(14)
    prelude_style.font.bold = True
    prelude_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    prelude_style.paragraph_format.first_line_indent = Pt(0)

    instruction_style = doc.styles.add_style("Instruction", 1)
    instruction_style.font.name = "Times New Roman"
    instruction_style.font.size = Pt(11)
    instruction_style.font.italic = True
    instruction_style.paragraph_format.first_line_indent = Pt(0)

    # ===== Cover page =====
    def add_centered(text: str, style_name: str | None = None, underline=False):
        para = doc.add_paragraph(text)
        if style_name:
            para.style = style_name
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.paragraph_format.first_line_indent = Pt(0)
        if underline:
            para.runs[0].underline = True
        return para

    add_centered("TRƯỜNG ĐẠI HỌC NGOẠI THƯƠNG", "PreludeHeading", underline=True)
    add_centered("KHOA KINH TẾ QUỐC TẾ", "PreludeHeading")
    doc.add_paragraph("")
    add_centered("KHÓA LUẬN TỐT NGHIỆP", "PreludeHeading")
    add_centered("[TÊN ĐỀ TÀI – VIẾT IN HOA, CĂN GIỮA]", "PreludeHeading")

    for label in (
        "Chuyên ngành",
        "Giảng viên hướng dẫn",
        "Sinh viên thực hiện",
        "Mã sinh viên",
    ):
        para = add_centered(f"{label}: ..............................................")
        para.paragraph_format.space_after = Pt(0)

    add_centered(f"[ĐỊA ĐIỂM], {date.today().year}")
    doc.add_page_break()

    # ===== Bìa phụ =====
    add_centered("BÌA PHỤ", "PreludeHeading")
    doc.add_paragraph(
        "Sao chép nguyên bản bìa chính trên giấy thường.", style="Instruction"
    )
    doc.add_page_break()

    # ===== Lời cảm ơn =====
    doc.add_paragraph("LỜI CẢM ƠN", style="ChapterTitle")
    doc.add_paragraph(
        "Viết tối đa 1 trang, đặt riêng trước Mục lục.",
        style="Instruction",
    )

    # ===== Mục lục & danh mục =====
    doc.add_paragraph("MỤC LỤC", style="ChapterTitle")
    doc.add_paragraph(
        "Dùng Insert → Table of Contents và cập nhật sau khi hoàn tất nội dung.",
        style="Instruction",
    )

    for title, note in (
        (
            "DANH MỤC CÁC CHỮ VIẾT TẮT",
            "Sắp xếp alphabet, ghi rõ nghĩa tiếng Việt và tiếng nước ngoài.",
        ),
        (
            "DANH MỤC CÁC BẢNG BIỂU",
            "Đánh số theo chương (ví dụ: Bảng 1.1); đặt tên phía trên bảng.",
        ),
        (
            "DANH MỤC CÁC HÌNH VẼ",
            "Đánh số theo chương (ví dụ: Hình 2.1); đặt tên và nguồn phía dưới.",
        ),
    ):
        doc.add_paragraph(title, style="ChapterTitle")
        doc.add_paragraph(note, style="Instruction")

    # ===== Tóm tắt =====
    doc.add_paragraph("TÓM TẮT KHÓA LUẬN", style="ChapterTitle")
    doc.add_paragraph(
        "200–250 từ: mục tiêu, phương pháp, dữ liệu, kết quả chính và đóng góp."
    )

    # ===== Phần mở đầu =====
    doc.add_paragraph("PHẦN MỞ ĐẦU", style="ChapterTitle")
    for item in (
        "1. Tính cấp thiết của nghiên cứu",
        "2. Mục tiêu chung và mục tiêu cụ thể",
        "3. Đối tượng và phạm vi nghiên cứu (nội dung, thời gian, không gian)",
        "4. Phương pháp nghiên cứu và dữ liệu sử dụng",
        "5. Cấu trúc của khóa luận",
    ):
        doc.add_paragraph(item)
    doc.add_paragraph(
        "Lưu ý: Không đặt Lời cảm ơn trong Phần mở đầu.", style="Instruction"
    )

    # ===== Chapters =====
    chapters = (
        (
            "CHƯƠNG 1",
            "TỔNG QUAN NGHIÊN CỨU VÀ/CƠ SỞ LÝ LUẬN",
            (
                "1.1 Tổng quan nghiên cứu quốc tế và trong nước",
                "1.2 Cơ sở lý thuyết và mô hình nghiên cứu",
                "1.3 Khoảng trống nghiên cứu và đóng góp dự kiến",
            ),
        ),
        (
            "CHƯƠNG 2",
            "DỮ LIỆU VÀ PHƯƠNG PHÁP NGHIÊN CỨU",
            (
                "2.1 Thiết kế nghiên cứu và mô hình",
                "2.2 Dữ liệu (GPR, giá tài sản, sự kiện)",
                "2.3 Phương pháp (Event Study, Wavelet QQR, ...)",
            ),
        ),
        (
            "CHƯƠNG 3",
            "KẾT QUẢ NGHIÊN CỨU",
            (
                "3.1 Phân tích phản ứng theo loại sự kiện",
                "3.2 So sánh phản ứng BTC – GOLD – OIL",
                "3.3 QQR và các kiểm định bổ sung",
            ),
        ),
        (
            "CHƯƠNG 4",
            "PHÂN TÍCH VÀ THẢO LUẬN",
            (
                "4.1 Tổng hợp các pattern chính",
                "4.2 Giải thích cơ chế theo vùng/kênh truyền dẫn",
                "4.3 Đối chiếu với nghiên cứu trước",
            ),
        ),
        (
            "CHƯƠNG 5",
            "HÀM Ý CHÍNH SÁCH / GIẢI PHÁP",
            (
                "5.1 Khuyến nghị cho nhà đầu tư",
                "5.2 Gợi ý cho nhà hoạch định chính sách",
                "5.3 Hạn chế nghiên cứu và hướng tiếp theo",
            ),
        ),
    )

    for code, title, items in chapters:
        doc.add_page_break()
        doc.add_paragraph(code, style="ChapterTitle")
        doc.add_paragraph(title, style="ChapterTitle")
        for item in items:
            doc.add_paragraph(item)

    # ===== Closing sections =====
    tail_sections = (
        ("KẾT LUẬN", "2–3 trang: đóng góp, hạn chế, hướng nghiên cứu tiếp theo."),
        (
            "DANH MỤC TÀI LIỆU THAM KHẢO",
            "Theo chuẩn Tạp chí Kinh tế & Quản lý (FTU).",
        ),
        ("PHỤ LỤC", "Đặt các bảng/phân tích bổ sung, script, dữ liệu chi tiết."),
    )
    for title, note in tail_sections:
        doc.add_page_break()
        doc.add_paragraph(title, style="ChapterTitle")
        doc.add_paragraph(note)

    doc.save(output_file)


if __name__ == "__main__":
    target_dir = Path("docs/thesis/templates")
    target_dir.mkdir(parents=True, exist_ok=True)
    out_file = target_dir / "KLTN_FORMAT_TEMPLATE.docx"
    build_template(out_file)
    print(f"[OK] Template saved to {out_file}")

